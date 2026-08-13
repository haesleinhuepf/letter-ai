from pathlib import Path

from docx import Document

from letter_ai.cli import read_command, write_command
from letter_ai.documents import extract_text, write_docx
from letter_ai.llm import LLMSettings


class FakeLLM:
    def __init__(self, responses: list[str]) -> None:
        self.responses = iter(responses)
        self.settings = LLMSettings()
        self.messages: list[list[dict[str, str]]] = []

    def complete(self, messages: list[dict[str, str]]) -> str:
        self.messages.append(messages)
        return next(self.responses)


def test_extract_text_omits_headers_and_footers(tmp_path: Path) -> None:
    path = tmp_path / "letter.docx"
    document = Document()
    section = document.sections[0]
    section.header.paragraphs[0].text = "Private header"
    section.footer.paragraphs[0].text = "Page 1"
    document.add_paragraph("Main letter text")
    document.save(path)

    assert extract_text(path) == "Main letter text"


def test_read_archives_classified_letter(tmp_path: Path) -> None:
    input_path = tmp_path / "letter.docx"
    document = Document()
    document.add_paragraph("I support this applicant.")
    document.save(input_path)
    llm = FakeLLM(['{"type":"support letter","summary":["Supports the applicant"]}'])

    read_command(input_path, tmp_path / "archive", llm)

    assert (tmp_path / "archive" / "supportletter1.txt").read_text(encoding="utf-8").strip() == "I support this applicant."
    assert (tmp_path / "archive" / "supportletter1_summary.txt").read_text(encoding="utf-8").strip() == "- Supports the applicant"


def test_read_archives_plain_text_letter(tmp_path: Path) -> None:
    input_path = tmp_path / "letter.txt"
    input_path.write_text("I support this applicant.", encoding="utf-8")
    llm = FakeLLM(['{"type":"support letter","summary":["Supports the applicant"]}'])

    read_command(input_path, tmp_path / "archive", llm)

    assert (tmp_path / "archive" / "supportletter1.txt").read_text(encoding="utf-8").strip() == "I support this applicant."


def test_read_archives_review_report(tmp_path: Path) -> None:
    input_path = tmp_path / "report.docx"
    document = Document()
    document.add_paragraph("The application needs revision.")
    document.save(input_path)
    llm = FakeLLM(['{"type":"review report","summary":["Needs revision"]}'])

    read_command(input_path, tmp_path / "archive", llm)

    assert (tmp_path / "archive" / "reviewreport1.txt").exists()
    assert (tmp_path / "archive" / "reviewreport1_summary.txt").read_text(encoding="utf-8").strip() == "- Needs revision"


def test_write_docx_replaces_placeholder_text_in_place(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    document = Document()
    document.add_paragraph("Before")
    document.add_paragraph("Enter content here")
    document.add_paragraph("After")
    document.save(template)

    output = tmp_path / "out.docx"
    write_docx("Hello world", output, template, "Jane Doe")

    paragraphs = [paragraph.text for paragraph in Document(output).paragraphs]
    assert paragraphs == ["Before", "Hello world", "After"]
    assert "Jane Doe" not in paragraphs


def test_write_docx_replaces_placeholder_in_table_cells(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    document = Document()
    document.add_paragraph("Before")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Intro Enter content here end"
    document.add_paragraph("After")
    document.save(template)

    output = tmp_path / "out.docx"
    write_docx("Hello world", output, template, "Jane Doe")

    doc = Document(output)
    assert [paragraph.text for paragraph in doc.paragraphs] == ["Before", "After"]
    assert doc.tables[0].cell(0, 0).text == "Intro Hello world end"
    assert "Jane Doe" not in [paragraph.text for paragraph in doc.paragraphs]


def test_write_uses_matching_examples_and_attribution(tmp_path: Path) -> None:
    archive = tmp_path / "archive"
    archive.mkdir()
    (archive / "supportletter1.txt").write_text("Example support letter", encoding="utf-8")
    (archive / "supportletter1_summary.txt").write_text("- Supports someone", encoding="utf-8")
    (archive / "recommendationletter1.txt").write_text("Wrong type example", encoding="utf-8")
    (archive / "recommendationletter1_summary.txt").write_text("- Wrong type", encoding="utf-8")
    request = tmp_path / "request.txt"
    request.write_text("Write a support letter for Alex.", encoding="utf-8")
    output = tmp_path / "out.docx"
    llm = FakeLLM([
        '{"type":"support letter","summary":["Request asks for support"]}',
        "Dear committee,\nI support Alex.",
    ])

    write_command(request, None, output, archive, llm)

    assert len(llm.messages[1]) == 4
    assert "Example support letter" in llm.messages[1][2]["content"]
    assert "Wrong type example" not in str(llm.messages[1])
    paragraphs = [paragraph.text for paragraph in Document(output).paragraphs]
    assert "Dear committee," in paragraphs
    assert any("github.com/haesleinhuepf/letter-ai" in paragraph for paragraph in paragraphs)
