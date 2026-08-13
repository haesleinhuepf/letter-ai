from __future__ import annotations

from pathlib import Path

from docx import Document


def extract_text(path: Path) -> str:
    """Extract the main document body, excluding headers, footers, and page fields."""
    document = Document(path)
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8").strip()


def _iter_text_containers(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph

    for section in document.sections:
        for header_footer in (section.header, section.footer):
            for paragraph in header_footer.paragraphs:
                yield paragraph

    def walk_table(table):
        for row in table.rows:
            for cell in row.cells:
                yield cell
                for nested in cell.tables:
                    yield from walk_table(nested)

    for table in document.tables:
        yield from walk_table(table)


def write_docx(text: str, output: Path, template: Path | None) -> None:
    document = Document(template) if template else Document()
    placeholder_text = "Enter content here"
    replaced_placeholder = False

    for container in _iter_text_containers(document):
        if placeholder_text in container.text:
            container.text = container.text.replace(placeholder_text, text)
            replaced_placeholder = True

    if replaced_placeholder:
        output.parent.mkdir(parents=True, exist_ok=True)
        document.save(output)
        return

    body = document._element.body
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)

    for paragraph in text.splitlines():
        if paragraph.strip():
            document.add_paragraph(paragraph.strip())
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
